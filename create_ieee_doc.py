#!/usr/bin/env python3
"""
Generate IEEE 2-column format document for ConvoSim project
"""

import subprocess
import sys

# Install python-docx
subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'auto')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

# Create document
doc = Document()

# Set up margins for IEEE format
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("ConvoSim: An Interactive Convolution Simulator")
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

# Authors (subtitle)
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author.add_run("Web Application for Signal Processing Education")
author_run.font.size = Pt(11)
author_run.font.italic = True

# Add space
doc.add_paragraph()

# Create 2-column layout using table
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.allow_autofit = False

# Set column widths for 2-column layout
for row in table.rows:
    for idx, cell in enumerate(row.cells):
        cell.width = Inches(3.25)

# Left column
left_cell = table.rows[0].cells[0]
left_cell.vertical_alignment = 1  # Top alignment

# Add content to left column
left_para = left_cell.paragraphs[0]
left_para.text = ""

# I. Abstract
heading = left_cell.add_paragraph()
heading_run = heading.add_run("I. ABSTRACT")
heading_run.font.size = Pt(12)
heading_run.font.bold = True
heading_run.font.color.rgb = RGBColor(0, 51, 102)

abstract_text = left_cell.add_paragraph(
    "ConvoSim is an interactive web-based visualization tool designed to demystify the complex "
    "mathematical operation of convolution. This application provides real-time, step-by-step visualization "
    "of signal processing operations using Plotly.js. Students can interactively manipulate input signals x(t) "
    "and impulse responses h(t) in both continuous and discrete domains, observing the effects of signal flipping, "
    "shifting, multiplication, and summation. The tool supports both frequency and time-domain analysis, making it "
    "an invaluable educational resource for digital signal processing courses."
)
abstract_text.paragraph_format.space_after = Pt(6)
abstract_text.paragraph_format.line_spacing = 1.15

# II. Introduction
heading = left_cell.add_paragraph()
heading.paragraph_format.space_before = Pt(6)
heading_run = heading.add_run("II. INTRODUCTION")
heading_run.font.size = Pt(12)
heading_run.font.bold = True
heading_run.font.color.rgb = RGBColor(0, 51, 102)

intro_text = left_cell.add_paragraph(
    "Convolution is a fundamental operation in signal processing, yet students often struggle to visualize "
    "and understand its mechanics. Traditional textbook approaches rely on static diagrams and manual calculations, "
    "which can be tedious and error-prone. ConvoSim bridges this educational gap by providing an interactive, "
    "responsive platform where learners can instantly see how signals interact, flip, shift, and combine to produce "
    "output signals."
)
intro_text.paragraph_format.space_after = Pt(6)
intro_text.paragraph_format.line_spacing = 1.15

# III. Key Features
heading = left_cell.add_paragraph()
heading.paragraph_format.space_before = Pt(6)
heading_run = heading.add_run("III. KEY FEATURES")
heading_run.font.size = Pt(12)
heading_run.font.bold = True
heading_run.font.color.rgb = RGBColor(0, 51, 102)

features = [
    "Dual-mode visualization (continuous & discrete signals)",
    "Real-time signal manipulation with interactive sliders",
    "Step-by-step animation of convolution process",
    "Frequency domain (FFT) analysis display",
    "Range selector for zooming into specific signal regions",
    "Multiple predefined signal templates (impulse, step, pulse, etc.)",
    "Responsive two-column layout for side-by-side comparisons",
    "Export capabilities for educational documentation"
]

for feature in features:
    feature_para = left_cell.add_paragraph(feature, style='List Bullet')
    feature_para.paragraph_format.space_after = Pt(3)
    feature_para.paragraph_format.line_spacing = 1.15

# Right column
right_cell = table.rows[0].cells[1]
right_cell.vertical_alignment = 1

# Add content to right column
right_para = right_cell.paragraphs[0]
right_para.text = ""

# IV. Technical Architecture
heading = right_cell.add_paragraph()
heading_run = heading.add_run("IV. TECHNICAL ARCHITECTURE")
heading_run.font.size = Pt(12)
heading_run.font.bold = True
heading_run.font.color.rgb = RGBColor(0, 51, 102)

tech_text = right_cell.add_paragraph(
    "ConvoSim is built on modern web technologies following the Model-View-Controller (MVC) pattern: "
)
tech_text.paragraph_format.space_after = Pt(3)
tech_text.paragraph_format.line_spacing = 1.15

tech_stack = [
    ("Frontend", "HTML5, CSS3 with Glass-morphism UI design, JavaScript for interactivity"),
    ("Visualization", "Plotly.js 2.32.0 for high-quality mathematical plotting"),
    ("Backend", "Node.js with Express.js framework for lightweight server operations"),
    ("Signal Processing", "WebGL-enabled canvas for efficient real-time computation"),
    ("Mathematics", "Custom DSP algorithms for FFT, convolution, and signal transformations")
]

for label, desc in tech_stack:
    para = right_cell.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.15
    bold_run = para.add_run(f"• {label}: ")
    bold_run.font.bold = True
    para.add_run(desc)

# V. Mathematical Foundation
heading = right_cell.add_paragraph()
heading.paragraph_format.space_before = Pt(6)
heading_run = heading.add_run("V. MATHEMATICAL FOUNDATION")
heading_run.font.size = Pt(12)
heading_run.font.bold = True
heading_run.font.color.rgb = RGBColor(0, 51, 102)

math_text = right_cell.add_paragraph()
math_text.paragraph_format.space_after = Pt(6)
math_text.paragraph_format.line_spacing = 1.15
math_text.add_run("Continuous Convolution: ")
math_text.add_run("y(t) = ∫ x(τ)h(t-τ)dτ").italic = True
math_text.add_run("\n\nDiscrete Convolution: ")
math_text.add_run("y[n] = Σ x[k]h[n-k]").italic = True

# VI. Educational Impact
heading = right_cell.add_paragraph()
heading.paragraph_format.space_before = Pt(6)
heading_run = heading.add_run("VI. EDUCATIONAL IMPACT")
heading_run.font.size = Pt(12)
heading_run.font.bold = True
heading_run.font.color.rgb = RGBColor(0, 51, 102)

impact_text = right_cell.add_paragraph(
    "ConvoSim enhances learning outcomes by providing immediate visual feedback, reducing cognitive load, "
    "and enabling hypothesis-driven exploration. Students develop deeper understanding through interactive discovery "
    "rather than passive consumption of static diagrams."
)
impact_text.paragraph_format.space_after = Pt(6)
impact_text.paragraph_format.line_spacing = 1.15

# VII. Conclusion & Future Work
heading = right_cell.add_paragraph()
heading.paragraph_format.space_before = Pt(6)
heading_run = heading.add_run("VII. CONCLUSION")
heading_run.font.size = Pt(12)
heading_run.font.bold = True
heading_run.font.color.rgb = RGBColor(0, 51, 102)

conclusion_text = right_cell.add_paragraph(
    "ConvoSim represents a modern approach to signal processing education, combining rigorous mathematics "
    "with intuitive visualization. Future enhancements include Laplace/Z-transform visualization, "
    "system stability analysis, and multi-signal convolution capabilities."
)
conclusion_text.paragraph_format.space_after = Pt(6)
conclusion_text.paragraph_format.line_spacing = 1.15

# Add footer with metadata
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run(
    "ConvoSim v1.0.0 | MIT License | Published: 2026 | "
    "Repository: github.com/signals_system_web"
)
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(128, 128, 128)
footer_run.italic = True

# Save document
output_path = "ConvoSim_IEEE_2Column_Presentation.docx"
doc.save(output_path)
print(f"✓ Document created successfully: {output_path}")
